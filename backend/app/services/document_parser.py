import io
import re
from pathlib import Path

ALLOWED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024
MIN_PDF_TEXT_CHARS = 400
MIN_PDF_FILE_BYTES = 80_000

MIME_BY_EXTENSION = {
    ".txt": "text/plain; charset=utf-8",
    ".md": "text/markdown; charset=utf-8",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".pdf": "application/pdf",
}

NOISE_LINE_RE = re.compile(r"^\d{1,3}$")


class UnsupportedDocumentError(ValueError):
    pass


def validate_extension(file_name: str) -> str:
    extension = Path(file_name).suffix.lower()
    if extension == ".doc":
        raise UnsupportedDocumentError(
            "Формат .doc не поддерживается. Сохраните документ Word как .docx и загрузите снова."
        )
    if extension not in ALLOWED_EXTENSIONS:
        raise UnsupportedDocumentError("Поддерживаются только файлы .txt, .md, .docx и .pdf")
    return extension


def guess_mime_type(extension: str) -> str:
    return MIME_BY_EXTENSION.get(extension, "application/octet-stream")


def _render_table(rows: list[list[str | None]]) -> str:
    prepared_rows: list[list[str]] = []
    for row in rows:
        cells = [str(cell or "").strip() for cell in row]
        if not any(cells):
            continue
        prepared_rows.append(cells)

    if not prepared_rows:
        return ""

    header_alpha = sum(sum(1 for ch in cell if ch.isalpha()) for cell in prepared_rows[0])
    has_text_header = header_alpha >= 2

    body_rows: list[list[str]] = [prepared_rows[0]]
    for row in prepared_rows[1:]:
        alpha_chars = sum(sum(1 for ch in cell if ch.isalpha()) for cell in row)
        if not has_text_header and alpha_chars < 2 and all(len(cell) <= 3 for cell in row):
            continue
        body_rows.append(row)

    lines = [" | ".join(body_rows[0])]
    for row in body_rows[1:]:
        lines.append(" | ".join(row))
    return "\n".join(lines)


def _clean_text_block(text: str) -> str:
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            continue
        if NOISE_LINE_RE.fullmatch(line):
            continue
        if len(line) <= 2 and not any(ch.isalpha() for ch in line):
            continue
        lines.append(line)
    return "\n".join(lines)


def _merge_text_variants(*variants: str) -> str:
    cleaned = [_clean_text_block(item) for item in variants if item]
    cleaned = [item for item in cleaned if item]
    if not cleaned:
        return ""
    cleaned.sort(key=len, reverse=True)
    return cleaned[0]


def _extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(file_bytes))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        table_text = _render_table([[cell.text for cell in row.cells] for row in table.rows])
        if table_text:
            blocks.append(table_text)

    text = "\n\n".join(blocks)
    if not text.strip():
        raise ValueError("Word-документ не содержит извлекаемого текста")
    return text


def _extract_page_with_pdfplumber(page) -> str:
    import pdfplumber

    variants = [
        page.extract_text(layout=True, x_tolerance=2, y_tolerance=3) or "",
        page.extract_text() or "",
    ]
    table_blocks: list[str] = []
    for table in page.extract_tables() or []:
        table_text = _render_table(table)
        if table_text:
            table_blocks.append(table_text)

    page_text = _merge_text_variants(*variants)
    if table_blocks:
        table_part = "\n\n".join(table_blocks)
        if page_text:
            return f"{page_text}\n\n{table_part}"
        return table_part
    return page_text


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    from pypdf import PdfReader

    blocks: list[str] = []
    pypdf_reader = PdfReader(io.BytesIO(file_bytes))

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for index, page in enumerate(pdf.pages):
            plumber_text = _extract_page_with_pdfplumber(page)
            pypdf_text = ""
            if index < len(pypdf_reader.pages):
                pypdf_text = pypdf_reader.pages[index].extract_text() or ""
            merged = _merge_text_variants(plumber_text, pypdf_text)
            if merged:
                blocks.append(f"[Страница {index + 1}]\n{merged}")

    text = "\n\n".join(blocks)
    if not text.strip():
        raise ValueError(
            "PDF-файл не содержит извлекаемого текста. "
            "Возможно, это скан без текстового слоя — загрузите текстовый PDF или DOCX."
        )
    if len(text) < MIN_PDF_TEXT_CHARS and len(file_bytes) >= MIN_PDF_FILE_BYTES:
        raise ValueError(
            "Из PDF извлечено слишком мало текста. "
            "Скорее всего, документ является сканом. Загрузите версию с текстовым слоем или DOCX."
        )
    return text


def extract_text(file_name: str, file_bytes: bytes) -> str:
    extension = validate_extension(file_name)
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ValueError("Размер файла не должен превышать 10 МБ")
    if not file_bytes.strip():
        raise ValueError("Файл пустой")

    if extension in {".txt", ".md"}:
        for encoding in ("utf-8", "utf-8-sig", "cp1251"):
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Не удалось прочитать текстовый файл. Используйте UTF-8.")
        if not text.strip():
            raise ValueError("Файл не содержит текста")
        return text

    if extension == ".docx":
        return _extract_text_from_docx(file_bytes)

    return _extract_text_from_pdf(file_bytes)
