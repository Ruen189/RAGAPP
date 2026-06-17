import io

import pytest
from docx import Document

from app.services.document_parser import UnsupportedDocumentError, _clean_text_block, extract_text, validate_extension


def test_validate_txt_md_docx_pdf():
    assert validate_extension("guide.txt") == ".txt"
    assert validate_extension("guide.md") == ".md"
    assert validate_extension("guide.docx") == ".docx"
    assert validate_extension("guide.pdf") == ".pdf"


def test_reject_doc_extension():
    with pytest.raises(UnsupportedDocumentError):
        validate_extension("legacy.doc")


def test_extract_txt():
    text = extract_text("notes.txt", "Scrum basics".encode("utf-8"))
    assert "Scrum" in text


def test_extract_md():
    text = extract_text("notes.md", "# Kanban\n\nFlow".encode("utf-8"))
    assert "Kanban" in text


def test_extract_docx_tables():
    document = Document()
    document.add_paragraph("Project metrics")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Sprint"
    table.cell(0, 1).text = "Velocity"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "24"

    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text("metrics.docx", buffer.getvalue())
    assert "Project metrics" in text
    assert "Sprint | Velocity" in text
    assert "1 | 24" in text


def test_clean_text_block_filters_isolated_page_numbers():
    cleaned = _clean_text_block("ГОСТ 7.32-2017\n\n5\n\n2\n\nТребования к оформлению отчёта")
    lines = cleaned.split("\n")
    assert "5" not in lines
    assert "2" not in lines
    assert "ГОСТ 7.32-2017" in cleaned
    assert "Требования к оформлению отчёта" in cleaned
